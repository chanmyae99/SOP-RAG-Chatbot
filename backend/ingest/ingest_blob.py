# ingest/ingest_blob.py
from dotenv import load_dotenv
load_dotenv()

import os
from io import BytesIO

import base64
from backend.ingest.image_extractor import extract_images_from_pdf
from backend.ingest.image_captioner import caption_image
from backend.services.blob_storage import upload_image


def make_safe_key(raw: str) -> str:
    return base64.urlsafe_b64encode(raw.encode("utf-8")).decode("utf-8")


from backend.services.blob_storage import (
    get_container_client,
    list_blobs,
    download_blob
)
from backend.services.azure_search import upload_documents

from openai import OpenAI
from pypdf import PdfReader
from docx import Document
import pandas as pd


# ==============================
# CONFIG
# ==============================
BLOB_CONNECTION_STRING = os.getenv("BLOB_CONNECTION_STRING")
CONTAINER_NAME = "AZURE_STORAGE_CONTAINER"
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

EMBEDDING_MODEL = "text-embedding-3-small"
CHUNK_SIZE = 500
OVERLAP = 50


# ==============================
# CLIENTS
# ==============================
openai_client = OpenAI(api_key=OPENAI_API_KEY)

container_client = get_container_client()



# ==============================
# READERS
# ==============================
def read_pdf(blob_bytes):
    reader = PdfReader(BytesIO(blob_bytes))
    pages = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text:
            pages.append({
                "text": text,
                "page_number": i + 1
            })

    return pages


# def read_docx(blob_bytes):
#     doc = Document(BytesIO(blob_bytes))
#     pages = []

#     for i, p in enumerate(doc.paragraphs):
#         if p.text.strip():
#             pages.append({
#                 "text": p.text,
#                 "page_number": f"Paragraph {i + 1}"
#             })

#     return pages

def read_docx(blob_bytes):
    #print("📘 DOCX reader activated")

    doc = Document(BytesIO(blob_bytes))
    results = []

    current_section = "Introduction"
    paragraph_counter = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            #print("🧩 Heading detected:", text)
            current_section = text
            paragraph_counter = 0
            continue

        paragraph_counter += 1

        results.append({
            "text": text,
            "page_number": None,
            "section": current_section,
            "paragraph_number": paragraph_counter
        })

    #print("📘 DOCX paragraphs parsed:", len(results))
    return results




# def read_xlsx(blob_bytes):
#     df = pd.read_excel(BytesIO(blob_bytes))
#     return [{
#         "text": df.to_string(index=False),
#         "page_number": None
#     }]

def read_xlsx(blob_bytes):
    xls = pd.ExcelFile(BytesIO(blob_bytes))
    results = []

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)

        for idx, row in df.iterrows():
            # Convert row into readable text
            row_text = ", ".join(
                f"{col}: {row[col]}"
                for col in df.columns
                if pd.notna(row[col])
            )

            results.append({
                "text": f"Sheet: {sheet_name} | Row {idx + 1} | {row_text}",
                "page_number": None,          # Excel has no pages
                "sheet_name": sheet_name,
                "row_number": idx + 1
            })

    return results


# ==============================
# CHUNKING
# ==============================
def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):
    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap

    return chunks


# ==============================
# EMBEDDING
# ==============================
def embed_text(text: str) -> list:
    response = openai_client.embeddings.create(
        model=EMBEDDING_MODEL,
        input=text
    )
    return response.data[0].embedding



# ==============================
# INGEST PIPELINE
# ==============================
def ingest_documents():
    documents = []

    blobs = list_blobs(container_client)

    for blob in blobs:
        print(f"📄 Processing {blob.name}")
        blob_bytes = download_blob(container_client, blob.name)

        if blob.name.endswith(".pdf"):
            pages = read_pdf(blob_bytes)
            images = extract_images_from_pdf(blob_bytes, source_name=blob.name)
            # after extracting images
            for img in images:
                caption = caption_image(img["image_bytes"])
                vector = embed_text(caption)

                blob_path = f"images/{blob.name}/page_{img['page_number']}/{img['file_name']}"
                upload_image(container_client, blob_path, img["image_bytes"])

                documents.append({
                    "metadata_storage_path": make_safe_key(f"{blob.name}|img|{img['file_name']}"),
                    "asset_type": "image",
                    "image_blob_path": blob_path,
                    "image_caption": caption,
                    "image_vector": vector,
                    "metadata_storage_name": blob.name,
                    "page_number": img["page_number"]})
        elif blob.name.endswith(".docx"):
            pages = read_docx(blob_bytes)
        elif blob.name.endswith(".xlsx"):
            pages = read_xlsx(blob_bytes)
        else:
            print(f"Skipped unsupported file: {blob.name}")
            continue
        
        

        for page in pages:
            # decide chunking strategy
            if blob.name.endswith(".pdf"):
                chunks = chunk_text(page["text"])
            else:
                # docx & xlsx: treat one unit as one chunk
                chunks = [page["text"]]
                

            for chunk_index, chunk in enumerate(chunks):
                page_no = page.get("page_number")

                # Stable key (DO NOT rely on page_number for docx/xlsx)
                if page_no is None:
                    raw_key = f"{blob.name}|chunk={chunk_index}"
                else:
                    raw_key = f"{blob.name}|page={page_no}|chunk={chunk_index}"
                safe_key = make_safe_key(raw_key)

                vector = embed_text(chunk)

                document = {
                    "metadata_storage_path": safe_key,
                    "asset_type": "text",
                    "content": chunk,
                    "content_vector": vector,
                    "metadata_storage_name": blob.name,
                    "page_number": page_no
                }

                # DOCX metadata
                if page.get("section"):
                    document["section"] = page["section"]
                if page.get("paragraph_number"):
                    document["paragraph_number"] = page["paragraph_number"]

                # XLSX metadata
                if page.get("sheet_name"):
                    document["sheet_name"] = page["sheet_name"]
                if page.get("row_number"):
                    document["row_number"] = page["row_number"]

                documents.append(document)





    print(f"Uploading {len(documents)} chunks to Azure AI Search")
    upload_documents(documents)


# ==============================
# ENTRY POINT
# ==============================
if __name__ == "__main__":
    ingest_documents()


